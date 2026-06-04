#!/usr/bin/env python3
"""
生成 AgenticSRE 指标3.9 测试方案文档（docx）及系统拓扑图（png）
运行: python3 测试/generate_test_doc.py
输出: 测试/测试大纲-课题四-中山大学-指标3.9.docx
      测试/系统拓扑图.png
"""

import os, sys
from pathlib import Path

# ── 生成系统拓扑图 ──────────────────────────────────────────
def generate_topology_diagram(output_path: str):
    """使用 matplotlib 生成 AgenticSRE 系统部署拓扑图"""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch

    plt.rcParams["font.sans-serif"] = ["Arial Unicode MS", "SimHei", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False

    W, H = 40, 28
    fig, ax = plt.subplots(1, 1, figsize=(W, H))
    ax.set_xlim(0, W)
    ax.set_ylim(0, H)
    ax.axis("off")
    ax.set_facecolor("white")
    fig.patch.set_facecolor("white")

    # ── 标题 ──
    ax.text(W / 2, H - 0.7, "AgenticSRE 智能运维平台 — 系统部署拓扑图",
            ha="center", va="center", fontsize=32, fontweight="bold", color="#1a1a2e")

    # ── 辅助函数 ──
    def box(x, y, w, h, title, color, subs=None, tsz=16, ssz=14, gap=0.55):
        p = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.15",
                           facecolor=color, edgecolor="#333", alpha=0.9, lw=1.8)
        ax.add_patch(p)
        ax.text(x + w / 2, y + h - 0.52, title,
                ha="center", va="center", fontsize=tsz, fontweight="bold", color="white")
        for i, s in enumerate(subs or []):
            ax.text(x + w / 2, y + h - 1.15 - i * gap, s,
                    ha="center", va="center", fontsize=ssz, color="#f0f0f0")

    def arrow(x1, y1, x2, y2, txt="", color="#555", lw=2.2, fsz=14):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="->,head_width=0.35,head_length=0.18",
                                    color=color, lw=lw))
        if txt:
            mx, my = (x1 + x2) / 2, (y1 + y2) / 2
            ax.text(mx + 0.05, my + 0.28, txt, ha="center", va="center",
                    fontsize=fsz, color=color, fontweight="bold")

    def region(x, y, w, h, txt, fsz=16):
        p = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.25",
                           facecolor="#f5f5f5", edgecolor="#aaa", alpha=0.5, lw=2, ls="--")
        ax.add_patch(p)
        ax.text(x + 0.4, y + h - 0.4, txt,
                ha="left", va="center", fontsize=fsz, fontstyle="italic",
                color="#333", fontweight="bold")

    # ════════════════════════════════════════════════════════
    # 第1层 — 客户端
    # ════════════════════════════════════════════════════════
    cy, ch = 24.8, 2.0
    region(0.5, cy, W - 1.0, ch, "客户端层", fsz=16)
    cw = 8.5
    cgap = (W - 1.0 - 4 * cw) / 5
    box(0.5 + cgap, cy + 0.25, cw, 1.5, "Web 浏览器", "#27ae60",
        ["Dashboard 交互界面"], tsz=16, ssz=14)
    box(0.5 + 2 * cgap + cw, cy + 0.25, cw, 1.5, "API 客户端", "#27ae60",
        ["REST API / SSE 实时推送"], tsz=16, ssz=14)
    box(0.5 + 3 * cgap + 2 * cw, cy + 0.25, cw, 1.5, "CLI 命令行", "#27ae60",
        ["main.py 多模式入口"], tsz=16, ssz=14)
    box(0.5 + 4 * cgap + 3 * cw, cy + 0.25, cw, 1.5, "MCP 客户端", "#27ae60",
        ["Claude / Copilot 集成"], tsz=16, ssz=14)

    # ════════════════════════════════════════════════════════
    # 第2层 — 跳板机
    # ════════════════════════════════════════════════════════
    jy, jh = 22.0, 2.2
    region(0.5, jy, W - 1.0, jh, "外部网络接入", fsz=16)
    box(10.0, jy + 0.3, 20.0, 1.6, "跳板机  (Jump Host)", "#e67e22",
        ["IP: 222.200.180.102   |   SSH 端口转发   |   用户: openstack"],
        tsz=18, ssz=16)

    # ════════════════════════════════════════════════════════
    # 第3层 — K8S 集群
    # ════════════════════════════════════════════════════════
    region(0.5, 0.5, W - 1.0, 21.0, "Kubernetes 集群  (入口节点: 10.10.3.110)", fsz=18)

    # ── 3a: 5个节点 ──
    ny, nh, nw = 17.2, 3.5, 7.2
    ng = (W - 1.0 - 5 * nw) / 6

    def nx(i):
        return 0.5 + ng + i * (nw + ng)

    box(nx(0), ny, nw, nh, "Master Node", "#2c3e50",
        ["k8s-master / 10.10.3.110", "Ubuntu 22.04  |  K8s v1.28",
         "8 vCPU / 32 GB / 500 GB",
         "apiserver, etcd, scheduler",
         "AgenticSRE Core 部署"],
        tsz=16, ssz=14, gap=0.50)

    box(nx(1), ny, nw, nh, "Worker Node 1", "#2c3e50",
        ["k8s-node1 / 10.10.3.111", "Ubuntu 22.04  |  K8s v1.28",
         "8 vCPU / 32 GB / 500 GB",
         "微服务应用 + 可观测后端",
         "Prometheus / Elasticsearch"],
        tsz=16, ssz=14, gap=0.50)

    box(nx(2), ny, nw, nh, "Worker Node 2", "#2c3e50",
        ["k8s-node2 / 10.10.3.112", "Ubuntu 22.04  |  K8s v1.28",
         "8 vCPU / 32 GB / 500 GB",
         "微服务应用 + 可观测后端",
         "Jaeger / Grafana"],
        tsz=16, ssz=14, gap=0.50)

    box(nx(3), ny, nw, nh, "GPU Node (Worker 3)", "#1a5276",
        ["gpu-node1 / 10.10.3.113", "Ubuntu 22.04  |  K8s v1.28",
         "8 vCPU / 64 GB / 1 TB",
         "NVIDIA A100 80 GB",
         "运维大模型 OpsLLM-14B"],
        tsz=16, ssz=14, gap=0.50)

    box(nx(4), ny, nw, nh, "KB Node (Worker 4)", "#1a5276",
        ["kb-node1 / 10.10.3.114", "Ubuntu 22.04  |  K8s v1.28",
         "8 vCPU / 32 GB / 1 TB",
         "Neo4j / ES / ChromaDB",
         "运维知识库服务节点"],
        tsz=16, ssz=14, gap=0.50)

    # ── 3b: AgenticSRE 平台组件 ──
    py, ph = 12.4, 3.5
    pw = 7.2
    pg = (W - 1.0 - 5 * pw) / 6

    def px(i):
        return 0.5 + pg + i * (pw + pg)

    box(px(0), py, pw, ph, "AgenticSRE Core", "#8e44ad",
        ["Web Dashboard (8080)", "Orchestrator (Pipeline/Daemon)",
         "MCP Server", "五阶段 Pipeline 编排引擎"],
        tsz=16, ssz=14, gap=0.50)

    box(px(1), py, pw, ph, "多智能体引擎", "#8e44ad",
        ["Alert / Metric / Log Agent", "Trace / Event / Hypothesis",
         "Correlation / Planning", "Remediation / Profiling"],
        tsz=16, ssz=14, gap=0.50)

    box(px(2), py, pw, ph, "记忆与演化层", "#8e44ad",
        ["FaultContextStore (ChromaDB)", "ContextLearner (规则学习)",
         "RCAJudge (质量评估)", "EvolutionTracker (演化追踪)"],
        tsz=16, ssz=14, gap=0.50)

    box(px(3), py, pw, ph, "运维大模型服务", "#c0392b",
        ["OpsLLM-14B (vLLM 推理)", "DeepSeek API 备用通道",
         "模型参数量 >= 14B", "对话 + API 双接入"],
        tsz=16, ssz=14, gap=0.50)

    box(px(4), py, pw, ph, "运维知识库", "#c0392b",
        ["故障处理知识库 (Neo4j)", "系统上下文知识库",
         "日志知识库 (ES)", "服务依赖 / 部署架构"],
        tsz=16, ssz=14, gap=0.50)

    # ── 3c: 工具层 + 可观测性 ──
    ty, th = 7.4, 3.5
    tw = 6.0
    tg = (W - 1.0 - 6 * tw) / 7

    def tx(i):
        return 0.5 + tg + i * (tw + tg)

    box(tx(0), ty, tw, th, "K8s 工具", "#d35400",
        ["kubectl 操作", "Pod / Deploy / Svc 管理",
         "事件监听 / dry-run", "安全操作 + 回滚"],
        tsz=16, ssz=14, gap=0.50)

    box(tx(1), ty, tw, th, "异常检测引擎", "#d35400",
        ["阈值检测 (Threshold)", "统计检测 (Z-score)",
         "趋势检测 (EWMA)", "多指标融合判定"],
        tsz=16, ssz=14, gap=0.50)

    box(tx(2), ty, tw, th, "Prometheus", "#e74c3c",
        ["指标采集存储 (:9090)", "node-exporter",
         "kube-state-metrics", "告警规则引擎"],
        tsz=16, ssz=14, gap=0.50)

    box(tx(3), ty, tw, th, "Elasticsearch", "#e74c3c",
        ["日志采集检索 (:9200)", "Filebeat 采集器",
         "全文检索 + 聚合", "日志知识存储"],
        tsz=16, ssz=14, gap=0.50)

    box(tx(4), ty, tw, th, "Jaeger", "#e74c3c",
        ["分布式追踪 (:16686)", "调用链采集",
         "Span 关联分析", "服务拓扑发现"],
        tsz=16, ssz=14, gap=0.50)

    box(tx(5), ty, tw, th, "Grafana", "#e74c3c",
        ["可视化监控 (:3000)", "运维仪表盘",
         "告警通知", "数据源聚合"],
        tsz=16, ssz=14, gap=0.50)

    # ── 3d: 被测应用 ──
    ay, ah = 3.2, 3.0
    box(1.0, ay, 17.5, ah,
        "被测微服务应用  (DeathStarBench Social Network)", "#16a085",
        ["nginx-thrift  |  compose-post  |  user-service  |  social-graph",
         "home-timeline  |  user-timeline  |  post-storage  |  media-service",
         "url-shorten  |  text-service"],
        tsz=16, ssz=14, gap=0.52)

    box(19.0, ay, 8.0, ah, "数据存储层", "#117a65",
        ["MongoDB (文档存储)", "Redis (缓存)",
         "Memcached (会话)"],
        tsz=16, ssz=14, gap=0.52)

    box(27.5, ay, 11.5, ah, "故障注入测试场景", "#7d3c98",
        ["OOMKilled / CrashLoopBackOff",
         "服务依赖故障 / 级联超时",
         "节点资源耗尽 / Pod 漂移"],
        tsz=16, ssz=14, gap=0.52)

    # ── 底部标注 ──
    ax.text(W / 2, 1.2,
            "中山大学  |  课题4: 基于大模型的云原生系统智能运维能力构建及演化  |  2026",
            ha="center", va="center", fontsize=16, color="#999", fontstyle="italic")

    # ════════════════════════════════════════════════════════
    # 连线
    # ════════════════════════════════════════════════════════
    mid = lambda x, w: x + w / 2

    # 客户端 → 跳板机
    for i in range(4):
        cx_i = 0.5 + cgap + i * (cw + cgap) + cw / 2
        arrow(cx_i, cy + 0.25, 20.0, jy + jh - 0.2, "")

    arrow(0.5 + cgap + cw / 2, cy + 0.25, 15.0, jy + jh - 0.2, "SSH 隧道", fsz=14)

    # 跳板机 → Master
    arrow(20.0, jy + 0.3, mid(nx(0), nw), ny + nh, "ssh -J", fsz=14)

    # Master → Workers
    for i in range(1, 5):
        arrow(nx(0) + nw, ny + nh / 2, nx(i), ny + nh / 2, "", color="#888")

    # 平台组件之间
    for i in range(4):
        arrow(px(i) + pw, py + ph / 2, px(i + 1), py + ph / 2, "", color="#8e44ad")

    # 平台 → 工具/可观测
    for i in range(5):
        arrow(px(i) + pw / 2, py, tx(i) + tw / 2 + 0.2, ty + th, "")

    # 工具 → 被测应用
    arrow(tx(0) + tw / 2, ty, 6.0, ay + ah, "管理", color="#888")
    arrow(tx(2) + tw / 2, ty, 10.0, ay + ah, "采集", color="#888")
    arrow(tx(3) + tw / 2, ty, 15.0, ay + ah, "采集", color="#888")
    arrow(tx(4) + tw / 2, ty, 22.0, ay + ah, "追踪", color="#888")

    plt.tight_layout(pad=0.3)
    fig.savefig(output_path, dpi=120, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    print(f"  拓扑图已生成: {output_path}")


# ── 生成 DOCX 文档 ──────────────────────────────────────────
def generate_test_docx(output_path: str, topology_img_path: str):
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()

    # ── 样式设置 ──
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def add_heading_cn(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = "黑体"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        return h

    def add_para(text, bold=False, alignment=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "宋体"
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        if bold:
            run.bold = True
        if alignment:
            p.alignment = alignment
        return p

    def set_cell_text(cell, text, bold=False, fontsize=9):
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = "宋体"
        run.font.size = Pt(fontsize)
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        if bold:
            run.bold = True

    def add_test_case_table(case_id, case_name, test_goal, preconditions, test_steps, input_data, expected_results, remarks="无"):
        """按模版格式添加测试用例表格"""
        table = doc.add_table(rows=7, cols=4, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Row 0: 用例编号 + 用例名称
        set_cell_text(table.cell(0, 0), "用例编号", bold=True)
        set_cell_text(table.cell(0, 1), case_id)
        set_cell_text(table.cell(0, 2), "用例名称", bold=True)
        set_cell_text(table.cell(0, 3), case_name)

        # Row 1: 测试目标 (合并后3列)
        set_cell_text(table.cell(1, 0), "测试目标", bold=True)
        table.cell(1, 1).merge(table.cell(1, 3))
        set_cell_text(table.cell(1, 1), test_goal)

        # Row 2: 前置条件 (合并后3列)
        set_cell_text(table.cell(2, 0), "前置条件", bold=True)
        table.cell(2, 1).merge(table.cell(2, 3))
        set_cell_text(table.cell(2, 1), preconditions)

        # Row 3: 测试过程 表头
        set_cell_text(table.cell(3, 0), "测试过程", bold=True)
        table.cell(3, 1).merge(table.cell(3, 2))
        set_cell_text(table.cell(3, 1), "步骤描述（页面操作/命令行操作）", bold=True)
        set_cell_text(table.cell(3, 3), "输入数据", bold=True)

        # Row 4: 测试过程 内容
        set_cell_text(table.cell(4, 0), "测试过程", bold=True)
        table.cell(4, 1).merge(table.cell(4, 2))
        set_cell_text(table.cell(4, 1), test_steps)
        set_cell_text(table.cell(4, 3), input_data)

        # Row 5: 预期结果 (合并后3列)
        set_cell_text(table.cell(5, 0), "预期结果", bold=True)
        table.cell(5, 1).merge(table.cell(5, 3))
        set_cell_text(table.cell(5, 1), expected_results)

        # Row 6: 备注
        set_cell_text(table.cell(6, 0), "备注", bold=True)
        table.cell(6, 1).merge(table.cell(6, 3))
        set_cell_text(table.cell(6, 1), remarks)

        doc.add_paragraph()  # 间距
        return table

    # ════════════════════════════════════════════════════════
    # 封面
    # ════════════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()
    add_para("国家重点研发计划", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para("课题测试方案", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(2):
        doc.add_paragraph()
    add_para("指标3.9：基于大模型的云原生软件系统智能化运维平台", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(4):
        doc.add_paragraph()
    add_para("中山大学", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para("2026年03月", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ════════════════════════════════════════════════════════
    # 1. 课题概述
    # ════════════════════════════════════════════════════════
    add_heading_cn("课题概述", level=1)

    # 课题信息表
    info_table = doc.add_table(rows=6, cols=4, style="Table Grid")
    info_data = [
        ["课题名称：", "基于大模型的云原生系统智能运维能力构建及演化", "", ""],
        ["所属项目：", "云原生软件生态系统智能化开发、测试与运维", "", ""],
        ["所属专项：", "先进计算与新兴软件", "", ""],
        ["项目牵头承担单位：", "", "", "阿里云计算有限公司"],
        ["课题承担单位：", "", "中山大学", ""],
        ["课题负责人：", "陈鹏飞", "", ""],
    ]
    for r, row_data in enumerate(info_data):
        for c, text in enumerate(row_data):
            if text:
                set_cell_text(info_table.cell(r, c), text, bold=(c == 0))
    # 合并
    for r in range(3):
        info_table.cell(r, 1).merge(info_table.cell(r, 3))
    info_table.cell(3, 0).merge(info_table.cell(3, 2))
    info_table.cell(4, 0).merge(info_table.cell(4, 1))
    doc.add_paragraph()

    # 1.1 测试目标
    add_heading_cn("测试目标", level=2)
    add_para(
        '本次测试的目的是对"云原生软件生态系统智能化开发、测试与运维"中课题4"基于大模型的云原生系统智能运维能力构建及演化"'
        "中指标3.9（基于大模型的云原生软件系统智能化运维原型平台）是否达到项目编号为2024YFB4505900的"
        "《国家重点研发计划项目任务书》中考核指标进行检测和确认。"
    )

    # 1.2 被测件列表
    add_heading_cn("被测件列表", level=2)
    bt = doc.add_table(rows=2, cols=6, style="Table Grid")
    for c, h in enumerate(["序号", "被测件名称", "版本", "运行方式", "单位", "备注"]):
        set_cell_text(bt.cell(0, c), h, bold=True)
    for c, v in enumerate(["1", "指标3.9：基于大模型的云原生软件系统智能化运维平台（AgenticSRE）", "V1.0", "B/S + CLI", "中山大学", "——"]):
        set_cell_text(bt.cell(1, c), v)
    doc.add_paragraph()

    # 1.3 技术创新性
    add_heading_cn("技术创新性", level=2)
    add_para(
        "本平台(AgenticSRE)针对云原生系统运维场景，创新性地提出了基于'发现-假设-规划-调查-推理'五阶段范式的多智能体协作架构。"
        "系统包含12个专用智能体(Alert/Metric/Log/Trace/Event/Detection/Hypothesis/Correlation/Planning/Profiling/"
        "Remediation/MetricAnomalyDetector)，支持链式(Chain)、反应式(ReAct)、反思式(Reflection)、规划执行式"
        "(Plan-and-Execute)、辩论式(Debate)、投票式(Voting)等6种多智能体协作范式。"
        "创新点包括：(1)多模态运维数据(指标/日志/调用链/告警/K8s事件)融合分析；"
        "(2)基于假设驱动的根因定位与交叉信号关联；"
        "(3)WeRCA式记忆学习与持续演化机制，支持专家反馈驱动的规则自动生成；"
        "(4)告警语义压缩与根因推荐；"
        "(5)安全的自愈操作与ActionStack回滚机制。"
    )

    # 1.4 对项目的贡献
    add_heading_cn("对项目的贡献", level=2)
    add_para(
        "本课题研制的AgenticSRE智能运维平台为项目提供了完整的云原生系统智能化运维解决方案，"
        "实现了从被动响应到主动诊断与自适应进化的跨越。平台通过多智能体协作自动形成智能化运维工作流，"
        "支持面向高可用性的运维逻辑与演化方案自动构造，有效提升了故障处理效率和运维自动化水平。"
    )

    # 1.5 课题概述及考核指标
    add_heading_cn("课题概述及考核指标", level=2)
    kt = doc.add_table(rows=2, cols=5, style="Table Grid")
    kt_headers = ["对应的课题名称", "指标名称", "立项时已有指标值/状态", "中期指标值/状态", "完成时指标值/状态"]
    for c, h in enumerate(kt_headers):
        set_cell_text(kt.cell(0, c), h, bold=True)
    kt_data = [
        "课题4：基于大模型的云原生系统智能运维能力构建及演化",
        "指标3.9：基于大模型的云原生软件系统智能化运维原型平台*",
        "当前智能运维平台缺少自动化运维逻辑构建及演化，故障处理效率低。",
        "完成基于大模型的云原生软件系统智能化运维原型平台的方案初步设计",
        "研制基于大模型的云原生软件系统智能化运维平台1套，自动形成智能化运维工作流，"
        "支持面向高可用性的运维逻辑与演化方案自动构造，故障处理时间比头部企业（如：阿里云）的故障处理时间减少10%",
    ]
    for c, v in enumerate(kt_data):
        set_cell_text(kt.cell(1, c), v)
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════
    # 2. 测试依据及引用文件
    # ════════════════════════════════════════════════════════
    add_heading_cn("测试依据及引用文件", level=1)
    add_heading_cn("测试依据", level=2)
    ref_t = doc.add_table(rows=4, cols=4, style="Table Grid")
    for c, h in enumerate(["序号", "名称", "发布日期", "编制单位"]):
        set_cell_text(ref_t.cell(0, c), h, bold=True)
    refs = [
        ["1", "国家重点研发计划项目任务书", "2024年12月17日", "中华人民共和国科学技术部制"],
        ["2", "GB/T 25000.51-2016《系统与软件工程 系统与软件质量要求和评价（SQuaRE）第51部分》",
         "2017年5月1日", "中华人民共和国国家质量监督检验检疫总局、中国国家标准化管理委员会"],
        ["3", "AgenticSRE系统设计说明", "2026年01月", "中山大学"],
    ]
    for r, row_data in enumerate(refs):
        for c, text in enumerate(row_data):
            set_cell_text(ref_t.cell(r + 1, c), text)
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════
    # 3. 测试环境
    # ════════════════════════════════════════════════════════
    add_heading_cn("测试环境", level=1)
    add_heading_cn("测试环境配置表", level=2)

    # 样品描述
    add_para("样品描述", bold=True)
    add_para(
        "本产品为基于大模型的云原生软件系统智能化运维平台(AgenticSRE)，"
        "采用'多智能体协作引擎+可观测性数据融合+记忆演化层'的总体架构。"
        "核心引擎包含12个专用智能体，基于'发现-假设-规划-调查-推理'五阶段Pipeline，"
        "支持链式、反应式、反思式、规划执行式、辩论式、投票式共6种协作范式。"
        "可观测性层融合Prometheus指标、Elasticsearch日志、Jaeger调用链等多模态数据。"
        "记忆演化层基于ChromaDB实现故障上下文存储与规则自动学习。"
        "平台支持Web界面、CLI、REST API和MCP Server四种接入方式，"
        "提供7×24异常监控与告警压缩、假设驱动的根因定位、"
        "运维工作流自动编排、高可用运维逻辑与演化方案自动构造等核心功能。"
    )
    doc.add_paragraph()

    # 检测环境表
    add_para("检测环境", bold=True)
    env_t = doc.add_table(rows=16, cols=4, style="Table Grid")
    env_t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    set_cell_text(env_t.cell(0, 0), "仪器编号", bold=True)
    env_t.cell(0, 1).merge(env_t.cell(0, 2))
    set_cell_text(env_t.cell(0, 1), "硬件环境", bold=True)
    set_cell_text(env_t.cell(0, 3), "软件环境", bold=True)

    set_cell_text(env_t.cell(1, 0), "仪器编号", bold=True)
    set_cell_text(env_t.cell(1, 1), "设备类型", bold=True)
    set_cell_text(env_t.cell(1, 2), "配置/性能参数", bold=True)
    set_cell_text(env_t.cell(1, 3), "软件环境", bold=True)

    # Master Node
    env_t.cell(2, 0).merge(env_t.cell(2, 3))
    set_cell_text(env_t.cell(2, 0), "管理节点 (Master Node)", bold=True)
    set_cell_text(env_t.cell(3, 0), "k8s-master\n10.10.3.110")
    set_cell_text(env_t.cell(3, 1), "服务器")
    set_cell_text(env_t.cell(3, 2),
                  "CPU型号：Intel(R) Xeon(R)\nCPU核数：8核\n内存规格：32GB\n存储规格：500GB SSD")
    set_cell_text(env_t.cell(3, 3),
                  "操作系统版本：Ubuntu 22.04 LTS\n"
                  "Kubernetes: v1.28\n"
                  "Docker/containerd: v1.7\n"
                  "Python: 3.10+\n"
                  "AgenticSRE: V1.0")

    # Worker Node 1
    env_t.cell(4, 0).merge(env_t.cell(4, 3))
    set_cell_text(env_t.cell(4, 0), "计算节点1 (Worker Node 1)", bold=True)
    set_cell_text(env_t.cell(5, 0), "k8s-node1\n10.10.3.111")
    set_cell_text(env_t.cell(5, 1), "服务器")
    set_cell_text(env_t.cell(5, 2),
                  "CPU型号：Intel(R) Xeon(R)\nCPU核数：8核\n内存规格：32GB\n存储规格：500GB SSD")
    set_cell_text(env_t.cell(5, 3),
                  "操作系统版本：Ubuntu 22.04 LTS\n"
                  "Kubernetes: v1.28\n"
                  "containerd: v1.7\n"
                  "Prometheus: v2.48\n"
                  "Elasticsearch: v8.11")

    # Worker Node 2
    env_t.cell(6, 0).merge(env_t.cell(6, 3))
    set_cell_text(env_t.cell(6, 0), "计算节点2 (Worker Node 2)", bold=True)
    set_cell_text(env_t.cell(7, 0), "k8s-node2\n10.10.3.112")
    set_cell_text(env_t.cell(7, 1), "服务器")
    set_cell_text(env_t.cell(7, 2),
                  "CPU型号：Intel(R) Xeon(R)\nCPU核数：8核\n内存规格：32GB\n存储规格：500GB SSD")
    set_cell_text(env_t.cell(7, 3),
                  "操作系统版本：Ubuntu 22.04 LTS\n"
                  "Kubernetes: v1.28\n"
                  "containerd: v1.7\n"
                  "Jaeger: v1.52\n"
                  "Grafana: v10.2")

    # Worker Node 3 - GPU节点 (运维大模型)
    env_t.cell(8, 0).merge(env_t.cell(8, 3))
    set_cell_text(env_t.cell(8, 0), "GPU计算节点 (Worker Node 3 - 运维大模型部署)", bold=True)
    set_cell_text(env_t.cell(9, 0), "gpu-node1\n10.10.3.113")
    set_cell_text(env_t.cell(9, 1), "GPU服务器")
    set_cell_text(env_t.cell(9, 2),
                  "CPU型号：Intel(R) Xeon(R)\nCPU核数：8核\n内存规格：64GB\n"
                  "存储规格：1TB SSD\nGPU：NVIDIA A100 80GB")
    set_cell_text(env_t.cell(9, 3),
                  "操作系统版本：Ubuntu 22.04 LTS\n"
                  "Kubernetes: v1.28\n"
                  "NVIDIA Driver: 535.x\n"
                  "CUDA: 12.2\n"
                  "vLLM: v0.4.x\n"
                  "运维大模型: OpsLLM-14B v3.0")

    # Worker Node 4 - 知识库节点
    env_t.cell(10, 0).merge(env_t.cell(10, 3))
    set_cell_text(env_t.cell(10, 0), "知识库节点 (Worker Node 4 - 运维知识库部署)", bold=True)
    set_cell_text(env_t.cell(11, 0), "kb-node1\n10.10.3.114")
    set_cell_text(env_t.cell(11, 1), "服务器")
    set_cell_text(env_t.cell(11, 2),
                  "CPU型号：Intel(R) Xeon(R)\nCPU核数：8核\n内存规格：32GB\n"
                  "存储规格：1TB SSD")
    set_cell_text(env_t.cell(11, 3),
                  "操作系统版本：Ubuntu 22.04 LTS\n"
                  "Kubernetes: v1.28\n"
                  "Neo4j: v5.14.0\n"
                  "Elasticsearch: v8.11.0\n"
                  "ChromaDB: v0.4.x\n"
                  "运维知识库: 故障处理/系统上下文/日志")

    # 客户端
    env_t.cell(12, 0).merge(env_t.cell(12, 3))
    set_cell_text(env_t.cell(12, 0), "客户端", bold=True)
    set_cell_text(env_t.cell(13, 0), "测试终端")
    set_cell_text(env_t.cell(13, 1), "笔记本电脑")
    set_cell_text(env_t.cell(13, 2), "CPU: Apple M1/Intel i7\n内存：16GB\n网络：千兆以太网")
    set_cell_text(env_t.cell(13, 3),
                  "macOS/Ubuntu\nChrome浏览器\nSSH客户端\nPython 3.10+")

    # 网络设备
    env_t.cell(14, 0).merge(env_t.cell(14, 3))
    set_cell_text(env_t.cell(14, 0), "网络接入", bold=True)
    set_cell_text(env_t.cell(15, 0), "跳板机\n222.200.180.102")
    set_cell_text(env_t.cell(15, 1), "跳板服务器")
    set_cell_text(env_t.cell(15, 2), "SSH端口转发\n外网IP: 222.200.180.102")
    set_cell_text(env_t.cell(15, 3),
                  "Ubuntu Server\nOpenSSH\n用户: openstack")
    doc.add_paragraph()

    # 拓扑结构图
    add_para("拓扑结构图", bold=True)
    if os.path.exists(topology_img_path):
        doc.add_picture(topology_img_path, width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para("图1 AgenticSRE智能运维平台系统部署拓扑图", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════
    # 4. 测试方法
    # ════════════════════════════════════════════════════════
    add_heading_cn("测试方法", level=1)
    add_para(
        "本次测试采用功能验证测试和性能对比测试相结合的方法。功能验证测试通过Web Dashboard界面操作和CLI命令行执行，"
        "验证平台各项功能是否满足设计要求；性能对比测试通过在真实K8S集群环境中注入典型故障场景，"
        "对比AgenticSRE平台与头部企业（阿里云）的故障处理时间，验证是否达到减少10%的考核指标。"
        "测试过程中，测试人员通过SSH跳板机（ssh -J openstack@222.200.180.102 ubuntu@10.10.3.110）"
        "远程登录K8S集群管理节点，执行相关测试操作并记录测试结果。"
    )

    # ════════════════════════════════════════════════════════
    # 5. 测试内容
    # ════════════════════════════════════════════════════════
    add_heading_cn("测试内容", level=1)

    # 5.1 考核指标与测试用例的对应关系
    add_heading_cn("考核指标与测试用例的对应关系", level=2)
    add_para("表 1 考核指标与测试用例的对应关系")
    map_t = doc.add_table(rows=6, cols=4, style="Table Grid")
    map_headers = ["序号", "技术指标", "测试项（考核指标）", "测试用例"]
    for c, h in enumerate(map_headers):
        set_cell_text(map_t.cell(0, c), h, bold=True)
    test_indicator = "指标3.9：基于大模型的云原生软件系统智能化运维原型平台*"
    test_requirement = (
        "研制基于大模型的云原生软件系统智能化运维平台1套，自动形成智能化运维工作流，"
        "支持面向高可用性的运维逻辑与演化方案自动构造，故障处理时间比头部企业（如：阿里云）的故障处理时间减少10%"
    )
    cases_list = [
        ("1", "智能化运维工作流自动形成能力测试"),
        ("2", "多智能体协作范式与故障诊断能力测试"),
        ("3", "面向高可用性的运维逻辑自动构造能力测试"),
        ("4", "运维演化方案自动构造与持续学习能力测试"),
        ("5", "故障处理时间对比测试（对比阿里云减少10%）"),
    ]
    for r, (num, name) in enumerate(cases_list):
        set_cell_text(map_t.cell(r + 1, 0), num)
        set_cell_text(map_t.cell(r + 1, 1), test_indicator)
        set_cell_text(map_t.cell(r + 1, 2), test_requirement)
        set_cell_text(map_t.cell(r + 1, 3), name)
    doc.add_paragraph()

    # 5.2 测试用例设计
    add_heading_cn("测试用例设计", level=2)
    add_para("5.2.1 基于大模型的云原生软件系统智能化运维平台用例设计", bold=True)

    # ── 测试用例 1: 智能化运维工作流自动形成 ──
    add_para("表 1 智能化运维工作流自动形成能力测试用例")
    add_test_case_table(
        case_id="PT-001",
        case_name="智能化运维工作流自动形成能力测试",
        test_goal=(
            "验证AgenticSRE平台是否能够根据输入的故障描述或告警信号，自动形成完整的智能化运维工作流，"
            "包括故障检测、假设生成、证据调查、根因推理和恢复建议五个阶段的自动编排与执行，"
            "无需人工干预即可完成端到端的故障诊断流程。"
        ),
        preconditions=(
            "1. AgenticSRE平台已在K8S集群（10.10.3.110）上完成部署，所有服务处于正常运行状态。\n"
            "2. 可观测性后端（Prometheus、Elasticsearch、Jaeger）已正常运行并采集数据。\n"
            "3. 被测微服务应用（DeathStarBench Social Network）已部署并正常运行。\n"
            "4. 测试人员已通过SSH跳板机（ssh -J openstack@222.200.180.102 ubuntu@10.10.3.110）登录管理节点。\n"
            "5. AgenticSRE Web Dashboard已启动（端口8080），可通过浏览器访问。\n"
            "6. LLM服务（DeepSeek API）连接正常。"
        ),
        test_steps=(
            "1. 通过SSH登录K8S管理节点，启动AgenticSRE Web Dashboard：\n"
            "   cd /path/to/AgenticSRE && python main.py web\n"
            "2. 在被测集群中注入一个典型故障（Pod CrashLoopBackOff）：\n"
            "   kubectl apply -f test_fault_crashloop.yaml -n social-network\n"
            "3. 方式一（Web界面）：打开浏览器访问Dashboard，在RCA分析页面输入故障描述，点击'开始分析'按钮。\n"
            "4. 方式二（CLI命令行）：执行Pipeline命令触发全流程分析。\n"
            "5. 观察系统是否自动执行五阶段Pipeline：\n"
            "   - Detection阶段：自动检测到Pod异常状态\n"
            "   - Hypothesis阶段：自动生成根因假设\n"
            "   - Investigation阶段：多智能体并行收集证据\n"
            "   - Reasoning阶段：交叉关联推理并输出根因报告\n"
            "   - Recovery阶段：给出修复建议\n"
            "6. 检查Pipeline执行日志，确认每个阶段均自动触发且无需人工干预。\n"
            "7. 记录Pipeline执行总时间、各阶段耗时及最终诊断结果。"
        ),
        input_data=(
            "输入数据1（Web界面）：\n"
            "故障描述：\"social-network命名空间下compose-post服务Pod频繁重启，"
            "日志中出现OOMKilled错误，请进行根因分析。\"\n\n"
            "输入数据2（CLI命令行）：\n"
            "python main.py pipeline \"Pod CrashLoopBackOff in namespace social-network, "
            "compose-post service shows OOMKilled\" -n social-network"
        ),
        expected_results=(
            "1. 系统能够根据故障描述自动形成完整的五阶段运维工作流（Detection→Hypothesis→Investigation→Reasoning→Recovery），"
            "全程无需人工干预。\n"
            "2. Detection阶段：系统自动检测到Pod异常、告警信号并完成告警压缩。\n"
            "3. Hypothesis阶段：系统自动生成至少1个与实际故障相关的根因假设。\n"
            "4. Investigation阶段：系统自动调度多个智能体（MetricAgent、LogAgent、EventAgent等）并行收集证据。\n"
            "5. Reasoning阶段：系统自动完成证据关联分析并输出结构化的根因诊断报告。\n"
            "6. Recovery阶段：系统自动给出可操作的修复建议。\n"
            "7. 整个Pipeline执行过程在Dashboard上实时展示进度，SSE推送正常。"
        ),
        remarks="Web Dashboard URL: http://10.10.3.110:8080\n通过SSH端口转发访问：ssh -L 8080:localhost:8080 -J openstack@222.200.180.102 ubuntu@10.10.3.110"
    )

    # ── 测试用例 2: 多智能体协作范式与故障诊断 ──
    add_para("表 2 多智能体协作范式与故障诊断能力测试用例")
    add_test_case_table(
        case_id="PT-002",
        case_name="多智能体协作范式与故障诊断能力测试",
        test_goal=(
            "验证AgenticSRE平台的多智能体协作框架能否支持多种协作范式（Chain/ReAct/Reflection/"
            "Plan-and-Execute/Debate/Voting），并通过不同范式对同一故障场景进行诊断，"
            "验证系统的多模态数据融合分析能力及根因定位准确性。"
        ),
        preconditions=(
            "1. AgenticSRE平台已部署并正常运行。\n"
            "2. 被测微服务应用已部署，可观测性后端正常采集数据。\n"
            "3. 在K8S集群中预先注入一个已知根因的故障场景（如：nginx-thrift服务CPU资源限制过低导致请求超时）。\n"
            "4. 测试人员已登录管理节点。\n"
            "5. 已确认所有6种协作范式模块可用（python main.py paradigm list）。"
        ),
        test_steps=(
            "1. 登录K8S管理节点，确认故障已注入且可观测数据正在采集。\n"
            "2. 依次使用6种协作范式执行故障诊断：\n"
            "   a) python main.py paradigm chain \"nginx-thrift request timeout, high latency\" -n social-network\n"
            "   b) python main.py paradigm react \"nginx-thrift request timeout, high latency\" -n social-network\n"
            "   c) python main.py paradigm reflection \"nginx-thrift request timeout, high latency\" -n social-network\n"
            "   d) python main.py paradigm plan_and_execute \"nginx-thrift request timeout, high latency\" -n social-network\n"
            "   e) python main.py paradigm debate \"nginx-thrift request timeout, high latency\" -n social-network\n"
            "   f) python main.py paradigm voting \"nginx-thrift request timeout, high latency\" -n social-network\n"
            "3. 记录每种范式的执行结果，包括：根因诊断结论、置信度评分、执行耗时、调用的智能体列表。\n"
            "4. 使用compare命令进行多范式对比分析：\n"
            "   python main.py compare --paradigms chain,react,reflection,plan_and_execute,debate,voting\n"
            "5. 验证系统是否融合了多模态数据（指标、日志、调用链、K8s事件）进行综合分析。\n"
            "6. 检查诊断报告中是否包含来自不同数据源的证据。"
        ),
        input_data=(
            "输入数据：\n"
            "故障描述：\"nginx-thrift request timeout, high latency observed in social-network namespace. "
            "Response time increased from 50ms to 2000ms. Multiple 504 errors in access logs.\"\n\n"
            "预设故障根因：nginx-thrift Deployment CPU资源限制(limits)设置为50m，"
            "在高并发下CPU被throttle导致请求超时。"
        ),
        expected_results=(
            "1. 所有6种协作范式均能成功执行完成，不出现系统错误或中断。\n"
            "2. 至少4种范式能够正确定位到CPU资源限制过低这一根因或给出密切相关的分析。\n"
            "3. 诊断报告中能够展示来自多模态数据源（Prometheus指标、ES日志、Jaeger调用链、K8s事件）的证据。\n"
            "4. compare对比结果能够清晰展示各范式的诊断准确率、置信度和耗时差异。\n"
            "5. 系统能够对12个专用智能体进行合理编排调度，各智能体职责分工明确。"
        ),
        remarks="6种范式说明：Chain(顺序链式)、ReAct(思考-行动-观察循环)、Reflection(自我反思改进)、"
                "Plan-and-Execute(先规划后执行)、Debate(多视角辩论)、Voting(多温度投票)"
    )

    # ── 测试用例 3: 高可用运维逻辑自动构造 ──
    add_para("表 3 面向高可用性的运维逻辑自动构造能力测试用例")
    add_test_case_table(
        case_id="PT-003",
        case_name="面向高可用性的运维逻辑自动构造能力测试",
        test_goal=(
            "验证AgenticSRE平台是否具备面向高可用性的运维逻辑自动构造能力，包括：（1）7x24持续异常监控与告警自动压缩；"
            "（2）多维度异常检测（阈值/Z-score/EWMA）自动识别；（3）告警去重与根因推荐；"
            "（4）基于安全策略的自愈操作建议生成与ActionStack回滚机制。"
        ),
        preconditions=(
            "1. AgenticSRE Daemon模式已配置并可启动（7x24持续监控模式）。\n"
            "2. Prometheus已配置node-exporter和kube-state-metrics，正常采集集群指标。\n"
            "3. 告警压缩功能已启用（config.yaml中alert.compression_enabled: true）。\n"
            "4. 被测微服务应用正常运行。\n"
            "5. 异常检测配置已就绪（threshold/zscore/ewma三种检测方法）。"
        ),
        test_steps=(
            "1. 启动AgenticSRE Daemon模式：\n"
            "   python main.py daemon -n social-network -i 30\n"
            "2. 等待系统完成首轮健康检查，确认Daemon正常运行。\n"
            "3. 在集群中注入多个并发故障以触发大量告警：\n"
            "   a) 对compose-post Pod施加CPU压力：kubectl exec -it compose-post-xxx -- stress --cpu 4\n"
            "   b) 人为制造mongodb连接池耗尽：模拟大量连接请求\n"
            "   c) 将home-timeline Pod副本数缩为0：kubectl scale deploy home-timeline --replicas=0\n"
            "4. 观察Daemon是否自动检测到上述异常并触发Pipeline。\n"
            "5. 检查告警压缩功能：\n"
            "   python main.py alert-scan -n social-network -r 15m\n"
            "   验证多个相关告警是否被正确压缩为告警组，并给出根因推荐。\n"
            "6. 检查系统生成的自愈建议是否包含安全检查项和回滚方案。\n"
            "7. 验证异常检测算法（阈值/Z-score/EWMA）是否对CPU、内存、磁盘等指标产生了正确的异常判定。\n"
            "8. 运行status命令查看系统整体健康状态：python main.py status"
        ),
        input_data=(
            "输入数据：\n"
            "Daemon配置参数：\n"
            "  poll_interval_seconds: 30\n"
            "  dedup_ttl_seconds: 300\n"
            "  max_concurrent_pipelines: 3\n"
            "  default_namespace: social-network\n\n"
            "检测方法配置：\n"
            "  node_cpu_usage: threshold(warn:85%, crit:95%) + zscore + ewma\n"
            "  node_memory_usage: threshold(warn:85%, crit:95%) + zscore\n"
            "  container_cpu: threshold(warn:80%, crit:95%) + zscore + ewma\n"
            "  container_memory: threshold(warn:80%, crit:95%)"
        ),
        expected_results=(
            "1. Daemon模式能够以30秒为周期持续轮询集群状态，自动检测到所有注入的异常。\n"
            "2. 告警压缩功能能够将相关联的多条告警正确聚合为告警组，压缩率≥50%。\n"
            "3. 每个告警组附带根因推荐说明，推荐准确率≥80%。\n"
            "4. 异常检测算法能够正确识别CPU使用率异常（Z-score越限）和Pod状态异常。\n"
            "5. 系统自动触发Pipeline对检测到的异常进行分析，无需人工介入。\n"
            "6. 生成的自愈建议包含明确的操作步骤、安全检查项及回滚预案（ActionStack）。\n"
            "7. 告警去重机制有效避免重复触发对同一故障的分析。"
        ),
        remarks="Daemon模式通过Ctrl+C停止。注入的故障需在测试完成后手动清理恢复。"
    )

    # ── 测试用例 4: 运维演化方案自动构造 ──
    add_para("表 4 运维演化方案自动构造与持续学习能力测试用例")
    add_test_case_table(
        case_id="PT-004",
        case_name="运维演化方案自动构造与持续学习能力测试",
        test_goal=(
            "验证AgenticSRE平台的持续演化能力，包括：（1）故障上下文的自动存储与知识积累；"
            "（2）基于历史故障的规则自动学习（ContextLearner）；（3）专家反馈驱动的监督学习机制；"
            "（4）RCA质量评估与自动优化；（5）演化趋势跟踪与报告。"
        ),
        preconditions=(
            "1. AgenticSRE平台已部署，记忆演化模块已启用（config.yaml中memory.enabled: true）。\n"
            "2. ChromaDB向量数据库服务正常运行。\n"
            "3. 已完成至少2次故障诊断流程，系统中已有历史故障记录。\n"
            "4. 演化追踪模块已启用（config.yaml中evolution.enabled: true）。"
        ),
        test_steps=(
            "1. 首先执行一次完整的故障诊断Pipeline：\n"
            "   python main.py pipeline \"MongoDB connection timeout in social-network\" -n social-network\n"
            "2. 检查故障上下文是否被自动存储：\n"
            "   查看./data/memory目录下ChromaDB中是否新增了故障记录。\n"
            "3. 再次执行类似的故障诊断，验证系统是否利用历史知识加速诊断：\n"
            "   python main.py pipeline \"MongoDB connection pool exhausted\" -n social-network\n"
            "   检查Hypothesis阶段是否注入了历史相似故障的知识。\n"
            "4. 提交专家反馈，触发监督学习：\n"
            "   python main.py feedback --incident-id <上次诊断的incident_id> "
            "--diagnosis \"Root cause is MongoDB connection pool size too small, "
            "need to increase maxPoolSize from 100 to 500\" --comment \"confirmed by DBA\"\n"
            "5. 验证专家反馈是否触发了新规则的自动生成。\n"
            "6. 查看演化报告：\n"
            "   python main.py evolution\n"
            "7. 验证演化报告中是否展示了知识库增长、诊断置信度变化、响应延迟趋势等信息。"
        ),
        input_data=(
            "输入数据1（故障诊断）：\n"
            "\"MongoDB connection timeout in social-network namespace, "
            "multiple services reporting database errors\"\n\n"
            "输入数据2（专家反馈）：\n"
            "--incident-id: 由首次诊断返回\n"
            "--diagnosis: \"Root cause is MongoDB connection pool size too small, "
            "need to increase maxPoolSize from 100 to 500\"\n"
            "--comment: \"confirmed by DBA team\""
        ),
        expected_results=(
            "1. 故障诊断完成后，故障上下文（包括告警信号、证据链、诊断结论）被自动存储到ChromaDB中。\n"
            "2. 第二次类似故障诊断时，系统Hypothesis阶段能够检索到历史相似故障并注入知识，诊断时间明显缩短。\n"
            "3. 专家反馈提交后，系统自动触发ContextLearner进行规则学习，返回结果显示rules_generated ≥ 1。\n"
            "4. 演化报告（python main.py evolution）正常输出，包含以下信息：\n"
            "   - 知识库规则增长数量（Net growth > 0）\n"
            "   - 诊断置信度趋势（trend为improving或stable）\n"
            "   - 响应延迟统计\n"
            "   - RCA质量评估分数\n"
            "5. 系统展现持续学习能力：随着故障案例和专家反馈的积累，诊断准确率和效率逐步提升。"
        ),
        remarks="记忆存储路径：./data/memory\n演化快照路径：./data/evolution\n"
                "ChromaDB backend配置：memory.backend: chromadb"
    )

    # ── 测试用例 5: 故障处理时间对比 ──
    add_para("表 5 故障处理时间对比测试用例（对比阿里云减少10%）")
    add_test_case_table(
        case_id="PT-005",
        case_name="故障处理时间对比测试（对比阿里云减少10%）",
        test_goal=(
            "验证AgenticSRE平台的故障处理时间是否比头部企业（阿里云）的故障处理时间减少10%。"
            "通过在相同K8S集群环境中注入多组典型故障场景，分别使用AgenticSRE平台和阿里云ARMS/SLS等运维工具进行故障诊断，"
            "记录从故障发生到给出根因诊断结论的完整时间（MTTD+MTTA），进行定量对比分析。"
        ),
        preconditions=(
            "1. AgenticSRE平台已部署并正常运行，Daemon模式可启动。\n"
            "2. 阿里云运维工具（ARMS应用监控/SLS日志服务/CloudMonitor）已配置为对比基准。\n"
            "3. 已准备至少3组典型故障场景的注入脚本。\n"
            "4. 计时工具就绪（精确到秒级）。\n"
            "5. 被测微服务应用处于正常负载运行状态。\n"
            "6. 已获取阿里云对同类故障场景的基准处理时间数据（来自公开技术报告或实测数据）。"
        ),
        test_steps=(
            "1. 准备故障场景注入脚本（3组典型故障）：\n"
            "   场景A：Pod OOMKilled（内存溢出导致Pod被杀）\n"
            "   场景B：Service依赖故障（下游服务不可用导致级联超时）\n"
            "   场景C：节点资源耗尽（Node NotReady导致Pod漂移）\n"
            "2. 对每组故障场景执行如下对比流程：\n"
            "   a) 启动AgenticSRE Daemon模式：python main.py daemon -n social-network\n"
            "   b) 记录故障注入时间T0。\n"
            "   c) 注入故障，启动计时。\n"
            "   d) 记录AgenticSRE检测到异常的时间T1（MTTD = T1-T0）。\n"
            "   e) 记录AgenticSRE输出根因诊断结论的时间T2（MTTA = T2-T1）。\n"
            "   f) 计算AgenticSRE总处理时间 = T2-T0。\n"
            "   g) 恢复系统，重新注入相同故障。\n"
            "   h) 使用阿里云ARMS/SLS工具进行故障诊断（或引用阿里云基准数据）。\n"
            "   i) 记录阿里云工具的总处理时间。\n"
            "3. 对3组场景的结果取平均值，计算时间减少百分比。\n"
            "4. 额外测试AgenticSRE在有历史知识积累情况下的加速效果：\n"
            "   重复注入场景A，验证第二次诊断是否更快。\n"
            "5. 生成对比分析报告。"
        ),
        input_data=(
            "场景A注入命令：\n"
            "kubectl set resources deploy compose-post -n social-network "
            "--limits=memory=10Mi\n"
            "（触发OOMKilled）\n\n"
            "场景B注入命令：\n"
            "kubectl scale deploy mongodb -n social-network --replicas=0\n"
            "（触发下游服务不可用）\n\n"
            "场景C注入命令：\n"
            "kubectl cordon k8s-node1 && kubectl drain k8s-node1 --ignore-daemonsets\n"
            "（触发节点不可用）\n\n"
            "阿里云基准数据来源：\n"
            "阿里云ARMS故障诊断基准时间（典型场景平均MTTR约15-30分钟，参考阿里云官方SRE白皮书）"
        ),
        expected_results=(
            "1. AgenticSRE平台在3组故障场景中均能自动完成从故障检测到根因诊断的全流程。\n"
            "2. 各场景处理时间记录如下（预期值）：\n"
            "   场景A（OOMKilled）：AgenticSRE ≤ 3分钟，阿里云基准约4-5分钟\n"
            "   场景B（依赖故障）：AgenticSRE ≤ 5分钟，阿里云基准约6-8分钟\n"
            "   场景C（节点故障）：AgenticSRE ≤ 4分钟，阿里云基准约5-7分钟\n"
            "3. AgenticSRE平均故障处理时间比阿里云基准减少≥10%。\n"
            "4. 有历史知识积累时，重复故障的诊断时间进一步缩短（预期减少20-30%）。\n"
            "5. 对比分析报告清晰展示各场景的时间对比数据、减少百分比及统计分析结果。\n"
            "6. 测试数据和计算过程透明可审计。"
        ),
        remarks=(
            "阿里云基准数据说明：\n"
            "1. 优先使用在相同集群环境中实测的阿里云ARMS/SLS工具诊断时间。\n"
            "2. 若无法实测，参考阿里云官方发布的SRE运维最佳实践报告中的基准MTTR数据。\n"
            "3. 时间统计口径统一为：从故障发生到输出根因诊断结论的时间（不含修复执行时间）。\n"
            "4. 每个场景至少执行3次取平均值以降低随机误差。"
        )
    )

    # ════════════════════════════════════════════════════════
    # 6. 项目管理、进度及人员分工
    # ════════════════════════════════════════════════════════
    add_heading_cn("项目管理、进度及人员分工", level=1)

    add_heading_cn("项目组织", level=2)
    add_para(
        "1) 实验室在接受该项目的测试委托任务后，将成立专门的测试项目组负责相关的测试工作。"
        "该项目组采用项目负责人负责制，项目负责人接受实验室管理部的质量监督和审查；"
    )
    add_para(
        "2) 项目负责人接受委托方及实验室管理部的监督，负责整个项目的测试过程进度、质量控制，"
        "负责协调安排测试资源，负责向委托方领导和实验室领导汇报测试情况；"
    )

    add_heading_cn("项目实施过程", level=1)
    steps = [
        "1) 项目组与委托单位进行详细的测试需求沟通，确定具体的测试需求；",
        "2) 项目组根据测试需求制定相应的测试方案和测试用例；",
        "3) 由项目负责人组织相关人员对测试方案和测试用例进行确认；",
        "4) 测试方案和测试用例确认后，项目组进行测试环境配置或确认工作；",
        "5) 测试环境确认后，项目组开始实施具体测试工作，并负责测试结果的确认工作，测试结束后项目组形成初步测试问题单；",
        "6) 项目组长组织质量监督人员及必要的技术人员对初步问题报告单进行审核，出现错误要求测试工程师进行重测或补测；",
        "7) 委托单位根据项目组提交的测试问题单进行被测软件的修改工作，并于40工作日前提交修改答复文档；",
        "8) 项目组对委托单位针对测试问题单的答复文档进行审查，并对修改后的产品进行回归测试，"
        "根据回归测试情况与委托方沟通后出具初步的测试报告，提交项目负责人和实验室授权签字人进行审核；",
        "9) 授权签字人审核批准后，项目组出具并提交产品最终测试报告。",
    ]
    for s in steps:
        add_para(s)

    add_heading_cn("质量保证", level=2)
    add_para(
        "在测试过程中，由项目负责人负责本项目测试过程中的测试质量监督和检查，"
        "并接受实验室管理部的指导，有权停止测试，并对出现的质量问题的纠正提供建议。"
    )

    add_heading_cn("测试计划进度", level=2)
    sched_t = doc.add_table(rows=8, cols=4, style="Table Grid")
    for c, h in enumerate(["测试活动", "责任人", "开始日期", "结束日期"]):
        set_cell_text(sched_t.cell(0, c), h, bold=True)
    sched_data = [
        ["填写测试派工单", "赵越鹏", "2026年3月20日", "2026年3月20日"],
        ["制定测试方案（测试计划）", "王丽娜", "2026年3月21日", "2026年3月22日"],
        ["制定测试记录", "王丽娜", "2026年3月23日", "2026年3月24日"],
        ["测试环境及设备确认", "王丽娜", "2026年3月25日", "2026年3月25日"],
        ["测试实施及记录", "王丽娜", "2026年3月26日", "2026年3月27日"],
        ["测试报告", "王丽娜", "2026年3月28日", "2026年3月28日"],
        ["归档", "王丽娜", "2026年3月31日", "2026年4月3日"],
    ]
    for r, row_data in enumerate(sched_data):
        for c, text in enumerate(row_data):
            set_cell_text(sched_t.cell(r + 1, c), text)
    doc.add_paragraph()

    add_heading_cn("人员分工", level=2)
    add_para(
        "本次标准符合性测试由中国电子技术标准化研究院赛西实验室组织，软件工程实验室作为承担该测试的测试单位，"
        "需要与委托方（课题组）发生业务交流和技术沟通，现将项目实施中各方需要给与的配合进行说明。"
        "项目负责人为赵越鹏，成员为王丽娜，工作内容如下表所示："
    )
    ppl_t = doc.add_table(rows=3, cols=4, style="Table Grid")
    for c, h in enumerate(["序号", "人员姓名", "角色", "工作内容"]):
        set_cell_text(ppl_t.cell(0, c), h, bold=True)
    ppl_data = [
        ["1", "赵越鹏", "项目负责人、质量监督员", "负责项目的整体控制与协调，负责对测试过程的质量进行监控"],
        ["2", "王丽娜", "测试人员", "按照项目负责人的安排，完成其对应的测试工作"],
    ]
    for r, row_data in enumerate(ppl_data):
        for c, text in enumerate(row_data):
            set_cell_text(ppl_t.cell(r + 1, c), text)
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════
    # 7. 测试终止条件
    # ════════════════════════════════════════════════════════
    add_heading_cn("测试终止条件", level=1)
    add_heading_cn("正常终止条件", level=2)
    add_para("本次测试拟开展一轮测试，若测试结束时：")
    add_para("已按测试合同或测试大纲要求完成所规定的测试；")
    add_para("则测试正常终止。")

    add_heading_cn("异常终止条件", level=2)
    add_para("若出现以下情况，则测试无法正常终止：")
    add_para("软件出现致命问题无法继续测试；")
    add_para("因委托方原因，无法完成测试。")

    # ════════════════════════════════════════════════════════
    # 8. 测试结果评价准则
    # ════════════════════════════════════════════════════════
    add_heading_cn("测试结果评价准则", level=1)
    add_heading_cn("测试结果评价准则", level=2)
    eval_t = doc.add_table(rows=2, cols=3, style="Table Grid")
    for c, h in enumerate(["测试内容", "评价结果类型", "说明"]):
        set_cell_text(eval_t.cell(0, c), h, bold=True)
    set_cell_text(eval_t.cell(1, 0), "功能性测试")
    set_cell_text(eval_t.cell(1, 1), "\"通过\"、\"不通过\"")
    set_cell_text(eval_t.cell(1, 2),
                  "测试用例中遇到的软件问题可分为两类：\n"
                  "1) 一般性问题：因出现错误导致测试过程无法正常完成，或者无法得到预期测试结果；\n"
                  "2) 建议性问题：测试过程能够正常完成，测试结果能够满足测试目标中的部分要求，但无法覆盖测试目标中的全部要求。\n"
                  "测试用例判定：\n"
                  "1) 当测试用例执行后，出现一般性问题和建议性问题时，此用例判定为测试不通过。\n"
                  "2) 当测试用例执行后，测试过程能够正常完成，且测试结果能够达到测试目标中的全部要求，此用例判定为测试通过。")
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════
    # 9. 输出成果
    # ════════════════════════════════════════════════════════
    add_heading_cn("输出成果", level=1)
    add_heading_cn("测试方案", level=2)
    add_para("软件测试方案指在正式测试实施开始前，对软件测试项目所作的一个执行方案，"
             "主要包括测试目的、评测依据、评测管理、评测内容及方法、测试配合要求、测试结果、测试环境要求以及项目输出成果等。")
    add_heading_cn("测试问题报告", level=2)
    add_para("测试问题报告指在测试实施完成后，测试工作组提交的一个测试问题报告。"
             "主要内容包括问题的严重等级、问题的结果描述等。")
    add_heading_cn("测试报告", level=2)
    add_para("测试报告是由测试工作组提交的最终测试结果报告，"
             "主要内容包括对软件产品的测试结论、详细测试结果描述以及软件的测试环境描述等。")

    # ════════════════════════════════════════════════════════
    # 10. 项目风险及控制
    # ════════════════════════════════════════════════════════
    add_heading_cn("项目风险及控制", level=1)
    add_para("本项目执行过程中，存在一定的项目进度风险。现将该风险及对应的风险控制手段说明如下：")

    add_heading_cn("风险因素", level=2)
    risks = [
        "系统的开发人员、技术人员不能及时提供现场支持，或者不能及时修改问题，导致项目延期；",
        "项目过程中出现技术瓶颈；",
        "项目过程周期预计出现偏差，导致项目延期；",
        "项目过程中出现突发事件，导致项目延期；",
        "测试过程中会产生一定的模拟数据，可能对目前系统造成部分影响，并且测试过程也有可能会影响对系统的正常使用；",
        "K8S集群远程访问（通过SSH跳板机）可能存在网络延迟或连接不稳定的风险。",
    ]
    for r in risks:
        add_para(r)

    add_heading_cn("风险控制及应对措施", level=2)
    measures = [
        "对于测试可能对系统产生的影响，建议将测试系统与真实业务系统进行隔离，并保证使用不同的数据库系统，并建议进行系统或者数据备份。",
        "在测试过程中，保证系统开发人员在现场提供必要的技术支持，必要时增派工程师，保证问题修改数量及质量。",
        "对于测试过程中的不可控因素及不可抗力，测试方和委托方本着友好协商的原则解决。",
        "对于远程访问风险，提前确认SSH跳板机和K8S集群网络连通性，准备备用接入方案。",
    ]
    for m in measures:
        add_para(m)

    # ── 保存 ──
    doc.save(output_path)
    print(f"✅ 测试文档已生成: {output_path}")


# ── Main ──────────────────────────────────────────────────
if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--skip-topology", action="store_true", help="跳过拓扑图生成，使用已有PNG")
    args = parser.parse_args()

    base_dir = Path(__file__).parent
    topology_path = str(base_dir / "系统拓扑图.png")
    docx_path = str(base_dir / "测试大纲-课题四-中山大学-指标3.9.docx")

    if not args.skip_topology:
        print("正在生成系统拓扑图...")
        generate_topology_diagram(topology_path)
    else:
        print(f"跳过拓扑图生成，使用已有文件: {topology_path}")

    print("正在生成测试方案文档...")
    generate_test_docx(docx_path, topology_path)

    print("\n✅ 所有文件生成完成！")
    print(f"  📄 测试文档: {docx_path}")
    print(f"  🖼️  拓扑图: {topology_path}")
